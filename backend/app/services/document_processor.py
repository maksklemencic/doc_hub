import io
import fitz
import re
import logging
from typing import List, Tuple
from PIL import Image
import pytesseract
from docx import Document
import unicodedata
import base64
import mimetypes

from ..errors.document_processor_errors import (
    DocumentProcessorError, UnsupportedDocumentTypeError, DocumentCorruptedError,
    EmptyDocumentError, TextExtractionError, OCRError, Base64DecodingError
)

logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    text = re.sub(r"-\n", "", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[\x00-\x1F\x7F]+", "", text)
    text = unicodedata.normalize("NFKC", text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> List[Tuple[int, str]]:
    """Extract text from PDF file using PyMuPDF and OCR fallback."""
    logger.info("Starting PDF text extraction")
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        logger.error(f"Failed to open PDF file: {str(e)}")
        raise DocumentCorruptedError("PDF", str(e))

    page_texts = []
    for i, page in enumerate(doc):
        page_number = i + 1
        try:
            text = page.get_text()

            if len(text.strip()) < 5:
                try:
                    logger.debug(f"Using OCR for page {page_number} due to insufficient text")
                    pix = page.get_pixmap(dpi=300, alpha=False)
                    img_bytes = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_bytes))
                    text = pytesseract.image_to_string(image)
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_number}: {str(e)}")
                    raise OCRError(f"OCR failed for page {page_number}: {str(e)}")
            
            cleaned_text = clean_text(text.strip())
            page_texts.append((page_number, cleaned_text))
        except Exception as e:
            if not isinstance(e, OCRError):
                logger.error(f"Text extraction failed for page {page_number}: {str(e)}")
                raise TextExtractionError("PDF", f"Page {page_number}: {str(e)}")
            raise
    
    if not page_texts or all(not text.strip() for _, text in page_texts):
        logger.warning("No readable text found in PDF")
        raise EmptyDocumentError("PDF")
    
    logger.info(f"Successfully extracted text from {len(page_texts)} pages")
    return page_texts


def extract_text_from_docx(file_bytes: bytes) -> List[Tuple[int, str]]:
    """Extract text from DOCX file including tables."""
    logger.info("Starting DOCX text extraction")
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"Failed to open DOCX file: {str(e)}")
        raise DocumentCorruptedError("DOCX", str(e))
    
    text_chunks = []
    for paragraph in doc.paragraphs:
        try:
            text = clean_text(paragraph.text)
            if text:
                text_chunks.append(text)
        except Exception as e:
            logger.warning(f"Failed to read paragraph: {str(e)}")

    try:
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(clean_text(cell.text) for cell in row.cells if cell.text.strip())
                if row_text:
                    text_chunks.append(row_text)
    except Exception as e:
        logger.warning(f"Failed to extract table content: {str(e)}")

    combined_text = "\n".join(text_chunks)

    if not combined_text.strip():
        logger.warning("No readable text found in DOCX file")
        raise EmptyDocumentError("DOCX")

    logger.info(f"Successfully extracted text from DOCX ({len(combined_text)} characters)")
    return [(1, combined_text)]


def extract_text_from_image(image_bytes: bytes) -> List[Tuple[int, str]]:
    """Extract text from image using OCR."""
    logger.info("Starting image OCR text extraction")
    try:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image).strip()
        clean = clean_text(text)
        
        if not clean.strip():
            logger.warning("No readable text found in image")
            raise EmptyDocumentError("image")
        
        logger.info(f"Successfully extracted text from image ({len(clean)} characters)")
        return [(1, clean)]
    except Exception as e:
        if not isinstance(e, EmptyDocumentError):
            logger.error(f"OCR failed for image: {str(e)}")
            raise OCRError(f"Image OCR failed: {str(e)}")
        raise


def normalize_file_type(mime_type: str) -> str:
    if mime_type.startswith("image/"):
        return "image"
    
    ext = mimetypes.guess_extension(mime_type)
    if not ext:
        logger.error(f"Unknown MIME type: {mime_type}")
        raise UnsupportedDocumentTypeError(mime_type, ["PDF", "DOCX", "images"])
    
    ext = ext.lstrip(".").lower()
    return ext


def process_document_for_text(file_bytes: bytes, mime_type: str) -> List[Tuple[int, str]]:
    """Process document and extract text based on MIME type."""
    logger.info(f"Processing document with MIME type: {mime_type}")
    
    handlers = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "img": extract_text_from_image
    }

    try:
        file_type = normalize_file_type(mime_type)
        handler = handlers.get(file_type)
        
        if not handler:
            logger.error(f"Unsupported document type: {file_type}")
            raise UnsupportedDocumentTypeError(file_type, list(handlers.keys()))
        
        result = handler(file_bytes)
        logger.info(f"Successfully processed {file_type} document")
        return result
    except (UnsupportedDocumentTypeError, DocumentCorruptedError, EmptyDocumentError, TextExtractionError, OCRError):
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing document: {str(e)}")
        raise DocumentProcessorError(f"Document processing failed: {str(e)}")

def base64_to_text(base64_text: str, mime_type: str) -> List[Tuple[int, str]]:
    """Decode base64 text and process as document."""
    logger.info("Processing base64 encoded document")
    try:
        file_bytes = base64.b64decode(base64_text)
        return process_document_for_text(file_bytes, mime_type)
    except Exception as e:
        if isinstance(e, (UnsupportedDocumentTypeError, DocumentCorruptedError, EmptyDocumentError, TextExtractionError, OCRError)):
            raise
        logger.error(f"Base64 decoding failed: {str(e)}")
        raise Base64DecodingError(str(e))